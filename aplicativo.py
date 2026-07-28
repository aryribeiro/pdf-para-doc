import io
import os
import re
import tempfile
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches, RGBColor
import docx.oxml as oxml
import docx.opc.constants as opc
from PIL import Image
from pypdf import PdfWriter
import pytesseract
import streamlit as st

# ==========================================
# Configurações da Página
# ==========================================
st.set_page_config(
    page_title="Conversor PDF p/ Docx", page_icon="📄", layout="centered"
)

# Logo do Web App e CSS/JS para Centralização
logo_url = "https://i.imgur.com/VNPhtmN.jpeg"
st.markdown(
    f"""
    <style>
        .centered-logo {{
            display: flex;
            justify-content: center;
            margin-bottom: 15px;
        }}
        div[data-testid="stRadio"] {{
            display: flex !important;
            flex-direction: column !important;
            align-items: center !important;
            justify-content: center !important;
            text-align: center !important;
            width: 100% !important;
        }}
        div[data-testid="stRadio"] > label,
        div[data-testid="stRadio"] [data-testid="stWidgetLabel"] {{
            text-align: center !important;
            width: 100% !important;
            display: flex !important;
            justify-content: center !important;
            white-space: nowrap !important;
        }}
        div[data-testid="stRadio"] [role="radiogroup"] {{
            display: inline-flex !important;
            flex-direction: column !important;
            align-items: flex-start !important;
            margin: 0 auto !important;
            width: max-content !important;
        }}
        div[data-testid="stRadio"] label p,
        div[data-testid="stRadio"] label span,
        div[data-testid="stRadio"] label {{
            white-space: nowrap !important;
            word-break: keep-all !important;
        }}
    </style>

    <script>
        function forcarLinhaUnicaECentralizar() {{
            const radioContainer = window.parent.document.querySelector('div[data-testid="stRadio"]');
            if (radioContainer) {{
                radioContainer.style.setProperty('display', 'flex', 'important');
                radioContainer.style.setProperty('flex-direction', 'column', 'important');
                radioContainer.style.setProperty('align-items', 'center', 'important');
                
                const labels = radioContainer.querySelectorAll('label, p, span');
                labels.forEach(el => {{
                    el.style.setProperty('white-space', 'nowrap', 'important');
                    el.style.setProperty('word-break', 'keep-all', 'important');
                }});
                
                const group = radioContainer.querySelector('div[role="radiogroup"]');
                if (group) {{
                    group.style.setProperty('margin', '0 auto', 'important');
                    group.style.setProperty('width', 'max-content', 'important');
                }}
            }}
        }}
        setTimeout(forcarLinhaUnicaECentralizar, 200);
        setTimeout(forcarLinhaUnicaECentralizar, 600);
    </script>

    <div class="centered-logo">
        <img src="{logo_url}" width="150">
    </div>
    """,
    unsafe_allow_html=True,
)


# Helper REESCRITO para aceitar parâmetros fiéis de fonte, tamanho e cor (Sem forçar Azul)
def adicionar_link_clicavel(paragraph, url, text, font_name="Arial", font_size_pt=11, hex_color="000000", is_bold=False, is_italic=False):
    part = paragraph.part
    r_id = part.relate_to(url, opc.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
    hyperlink = oxml.parse_xml(
        f'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        f'r:id="{r_id}"/>'
    )
    new_run = oxml.parse_xml(
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    )
    rPr = oxml.parse_xml(
        '<w:rPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    )
    
    if font_name:
        rFont = oxml.parse_xml(f'<w:rFonts xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>')
        rPr.append(rFont)
    if font_size_pt:
        sz = int(font_size_pt * 2) # Word usa meio-pontos
        rSz = oxml.parse_xml(f'<w:sz xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="{sz}"/>')
        rPr.append(rSz)
    if hex_color:
        c = oxml.parse_xml(f'<w:color xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="{hex_color}"/>')
        rPr.append(c)
    if is_bold:
        b = oxml.parse_xml('<w:b xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        rPr.append(b)
    if is_italic:
        i = oxml.parse_xml('<w:i xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        rPr.append(i)
        
    new_run.append(rPr)
    
    text_node = oxml.OxmlElement('w:t')
    text_node.set(oxml.ns.qn('xml:space'), 'preserve') # Preserva espaços com perfeição
    text_node.text = text
    new_run.append(text_node)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# ==========================================
# MÓDULO 1: Cópia Fiel (Para Impressão) -> MANTIDO INTACTO
# ==========================================
def modo_copia_fiel(pdf_file):
    pdf_bytes = pdf_file.read()
    doc_original = fitz.open(stream=pdf_bytes, filetype="pdf")
    merger = PdfWriter()

    for page_num in range(len(doc_original)):
        page = doc_original[page_num]
        pix = page.get_pixmap(dpi=300)
        img = Image.open(io.BytesIO(pix.tobytes()))

        pdf_ocr_bytes = pytesseract.image_to_pdf_or_hocr(
            img, extension="pdf", lang="por"
        )
        merger.append(io.BytesIO(pdf_ocr_bytes))

    with tempfile.NamedTemporaryFile(
        delete=False, suffix="_ocr.pdf"
    ) as temp_ocr_pdf:
        merger.write(temp_ocr_pdf)
        temp_ocr_pdf_path = temp_ocr_pdf.name

    temp_docx_path = temp_ocr_pdf_path.replace(".pdf", ".docx")

    try:
        from pdf2docx import Converter
        cv = Converter(temp_ocr_pdf_path)
        cv.convert(temp_docx_path, start=0, end=None)
        cv.close()

        with open(temp_docx_path, "rb") as f:
            return f.read()
    finally:
        if os.path.exists(temp_ocr_pdf_path):
            os.remove(temp_ocr_pdf_path)
        if os.path.exists(temp_docx_path):
            os.remove(temp_docx_path)


# ==========================================
# MÓDULO 2: Texto Editável (Limpeza de Artefatos) -> MANTIDO INTACTO
# ==========================================
def limpar_linha(texto_linha):
    if not texto_linha:
        return None

    l_lower = texto_linha.lower().strip()
    blacklist = ["firefox", "about:blank", "lofl", "1 of 1"]
    if any(termo in l_lower for termo in blacklist):
        return None

    texto_limpo = re.sub(r"^\s*[\(\[\{]?\d+[\)\]\}]?\s+(?=[A-Za-zÀ-ÿ])", "", texto_linha)
    texto_limpo = re.sub(r"^\s*[\(\[\{]?[A-Za-z0-9]{1,2}[\)\]\}]?\s+(?=[A-Za-zÀ-ÿ])", "", texto_limpo)
    texto_limpo = re.sub(r"^\s*[^a-zA-Z0-9À-ÿ]+\s*(?=[A-Za-zÀ-ÿ])", "", texto_limpo)

    return texto_limpo.strip() if texto_limpo.strip() else None


def modo_texto_editavel(pdf_file):
    doc_word = Document()
    pdf_bytes = pdf_file.read()
    pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")

    for page_num in range(len(pdf_document)):
        page = pdf_document[page_num]
        pix = page.get_pixmap(dpi=300)
        img = Image.open(io.BytesIO(pix.tobytes()))

        data = pytesseract.image_to_data(
            img, lang="por", output_type=pytesseract.Output.DICT
        )

        linhas = {}
        num_items = len(data["text"])

        for i in range(num_items):
            texto = data["text"][i].strip()
            confianca = int(data["conf"][i])

            if confianca > 30 and texto:
                block_num = data["block_num"][i]
                line_num = data["line_num"][i]
                chave_linha = (block_num, line_num)
                top = data["top"][i]
                height = data["height"][i]

                if chave_linha not in linhas:
                    linhas[chave_linha] = {"palavras": [], "top": top, "heights": []}

                linhas[chave_linha]["palavras"].append(texto)
                linhas[chave_linha]["heights"].append(height)

        linhas_ordenadas = sorted(linhas.values(), key=lambda x: x["top"])

        for item in linhas_ordenadas:
            texto_bruto = " ".join(item["palavras"])
            texto_linha = limpar_linha(texto_bruto)

            if not texto_linha:
                continue

            altura_media_px = sum(item["heights"]) / len(item["heights"])
            tamanho_fonte_pt = max(9, min(26, int(altura_media_px / 3.8)))

            p = doc_word.add_paragraph()
            run = p.add_run(texto_linha)
            run.font.name = "Arial"
            run.font.size = Pt(tamanho_fonte_pt)

            if tamanho_fonte_pt >= 14 or texto_linha.isdigit() or "Protocolo" in texto_linha:
                run.bold = True

            p.paragraph_format.space_after = Pt(4)

        if page_num < len(pdf_document) - 1:
            doc_word.add_page_break()

    docx_buffer = io.BytesIO()
    doc_word.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer.read()


# ==========================================
# MÓDULO 3: Fiel e Editável 100% NATIVO -> REESCRITO COM ALTA PRECISÃO
# ==========================================
def modo_fiel_editavel(pdf_file):
    pdf_bytes = pdf_file.read()
    doc_word = Document()
    pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    
    # Define Margem de segurança. Os cálculos posicionais compensarão a margem.
    MARGIN_PT = 36.0 
    for section in doc_word.sections:
        section.top_margin = Pt(MARGIN_PT)
        section.bottom_margin = Pt(MARGIN_PT)
        section.left_margin = Pt(MARGIN_PT)
        section.right_margin = Pt(MARGIN_PT)

    for page_idx in range(len(pdf_doc)):
        page = pdf_doc[page_idx]
        page_width = page.rect.width
        links_pagina = page.get_links()
        
        # 1. Extração Cirúrgica de Linhas Vetoriais (Pretas, Coloridas, etc)
        linhas_vetoriais = []
        for d in page.get_drawings():
            for item in d.get("items", []):
                if item[0] in ("l", "r"):
                    rect = item[1] if item[0] == "r" else fitz.Rect(item[1], item[2])
                    # Verifica se é uma linha horizontal
                    if rect.width > 20 and rect.height < 10:
                        color = d.get("color") or d.get("fill")
                        rgb_hex = "000000"
                        if color and len(color) == 3:
                            r, g, b = [int(c * 255) for c in color]
                            rgb_hex = f"{r:02x}{g:02x}{b:02x}"
                        
                        linhas_vetoriais.append({
                            "y0": rect.y0, 
                            "x0": rect.x0, 
                            "x1": rect.x1, 
                            "color": rgb_hex
                        })
        
        linhas_vetoriais.sort(key=lambda x: x["y0"])

        # 2. Extração de Texto com Coordenadas Preservadas
        blocks = page.get_text("dict", flags=fitz.TEXT_DEHYPHENATE)["blocks"]
        blocks = [b for b in blocks if b.get("type") == 0]
        blocks.sort(key=lambda x: x["bbox"][1])

        ultimo_y1 = MARGIN_PT
        linha_idx = 0

        for b in blocks:
            # 3. Inserir Linhas Vetoriais que aparecem antes do Bloco de Texto
            while linha_idx < len(linhas_vetoriais) and linhas_vetoriais[linha_idx]["y0"] < b["bbox"][1]:
                l_data = linhas_vetoriais[linha_idx]
                espaco_linha = max(0, l_data["y0"] - ultimo_y1)

                p_linha = doc_word.add_paragraph()
                p_linha.paragraph_format.space_before = Pt(espaco_linha)
                p_linha.paragraph_format.space_after = Pt(2)

                # Injecao XML para borda de parágrafo simulando vetor no exato layout
                pPr = p_linha._p.get_or_add_pPr()
                pbdr = oxml.parse_xml(
                    f'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">\n'
                    f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="{l_data["color"]}"/>\n'
                    f'</w:pBdr>'
                )
                pPr.append(pbdr)
                
                # Respeita o comprimento da linha calculando via Twips (20 por ponto)
                ind_left = int(max(0, l_data["x0"] - MARGIN_PT) * 20)
                ind_right = int(max(0, page_width - l_data["x1"] - MARGIN_PT) * 20)
                ind = oxml.parse_xml(
                    f'<w:ind xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                    f'w:left="{ind_left}" w:right="{ind_right}"/>'
                )
                pPr.append(ind)
                
                ultimo_y1 = l_data["y0"] + 2
                linha_idx += 1

            # 4. Inserir Bloco de Texto Perfeitamente Alinhado
            espaco_antes = max(0, b["bbox"][1] - ultimo_y1)
            p = doc_word.add_paragraph()
            p.paragraph_format.space_before = Pt(espaco_antes)
            p.paragraph_format.space_after = Pt(0) # Mantem linhas unidas como no PDF
            
            # Recuo Esquerdo Baseado no Eixo X Nativo do PDF
            indent_left = max(0, b["bbox"][0] - MARGIN_PT)
            p.paragraph_format.left_indent = Pt(indent_left)

            for l_idx, l in enumerate(b["lines"]):
                # Quebra de Linha: Nunca junta textos onde havia quebra no PDF
                if l_idx > 0:
                    run = p.add_run()
                    run.add_break()

                for span in l["spans"]:
                    texto = span["text"]
                    if not texto.strip() and texto != " ":
                        continue

                    # Extração de Atributos Nativos do Texto
                    font_size = span["size"]
                    font_name_clean = span["font"].split('+')[-1].split('-')[0]
                    if not font_name_clean: 
                        font_name_clean = "Arial"

                    e_bold = any(k in span["font"].lower() for k in ["bold", "black", "heavy"])
                    e_italic = any(k in span["font"].lower() for k in ["italic", "oblique"])

                    # Extração de Cor Hexadecimal Perfeita
                    c_val = span["color"]
                    r = (c_val >> 16) & 255
                    g = (c_val >> 8) & 255
                    b_col = c_val & 255
                    hex_color = f"{r:02x}{g:02x}{b_col:02x}"

                    # Validação de Hyperlinks no Span atual
                    uri_link = None
                    span_rect = fitz.Rect(span["bbox"])
                    for link in links_pagina:
                        if link.get("page") == page_idx or "uri" in link:
                            if span_rect.intersects(link["from"]) and "uri" in link:
                                uri_link = link["uri"]
                                break

                    # Inserção do Texto (Sendo Link ou Normal) preservando estilos Nativos
                    if uri_link:
                        adicionar_link_clicavel(
                            p, uri_link, texto,
                            font_name=font_name_clean,
                            font_size_pt=font_size,
                            hex_color=hex_color,
                            is_bold=e_bold,
                            is_italic=e_italic
                        )
                    else:
                        run = p.add_run(texto)
                        run.font.name = font_name_clean
                        run.font.size = Pt(font_size)
                        run.font.color.rgb = RGBColor(r, g, b_col)
                        run.bold = e_bold
                        run.italic = e_italic

            ultimo_y1 = b["bbox"][3]

        # 5. Inserir Restante de Linhas Vetoriais (Fim da Página)
        while linha_idx < len(linhas_vetoriais):
            l_data = linhas_vetoriais[linha_idx]
            espaco_linha = max(0, l_data["y0"] - ultimo_y1)

            p_linha = doc_word.add_paragraph()
            p_linha.paragraph_format.space_before = Pt(espaco_linha)
            p_linha.paragraph_format.space_after = Pt(2)

            pPr = p_linha._p.get_or_add_pPr()
            pbdr = oxml.parse_xml(
                f'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">\n'
                f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="{l_data["color"]}"/>\n'
                f'</w:pBdr>'
            )
            pPr.append(pbdr)
            
            ind_left = int(max(0, l_data["x0"] - MARGIN_PT) * 20)
            ind_right = int(max(0, page_width - l_data["x1"] - MARGIN_PT) * 20)
            ind = oxml.parse_xml(
                f'<w:ind xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                f'w:left="{ind_left}" w:right="{ind_right}"/>'
            )
            pPr.append(ind)
            linha_idx += 1

        if page_idx < len(pdf_doc) - 1:
            doc_word.add_page_break()

    docx_buffer = io.BytesIO()
    doc_word.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer.read()


# ==========================================
# Interface do Usuário (Streamlit UI)
# ==========================================
def main():
    st.markdown(
        "<h2 style='text-align: center; font-size: 1.6rem; margin-top: 0;'>Conversor PDF p/ Docx</h2>",
        unsafe_allow_html=True,
    )
    st.markdown(
        "<p style='text-align: center; color: #555555; margin-bottom: 20px; font-size: 0.95rem;'>"
        "Envie PDF para converter em Docx"
        "</p>",
        unsafe_allow_html=True,
    )

    col1, col2, col3 = st.columns([1, 2, 1])

    with col2:
        modo = st.radio(
            "Escolha o modo de conversão:",
            options=[
                "Cópia Fiel (Para Impressão: layout idêntico)",
                "Texto Editável (Texto limpo sem imagens)",
                "Fiel e Editável (Layout original e Textos Nativos)"
            ],
            index=2,
        )

        pdf_file = st.file_uploader(
            " ",
            type="pdf",
            label_visibility="collapsed",
        )

        if pdf_file:
            with st.spinner("Analisando estrutura e convertendo documento..."):
                try:
                    pdf_file.seek(0)

                    if "Cópia Fiel (Para Impressão" in modo:
                        docx_bytes = modo_copia_fiel(pdf_file)
                        nome_sufixo = "copia_fiel"
                    elif "Texto Editável" in modo:
                        docx_bytes = modo_texto_editavel(pdf_file)
                        nome_sufixo = "editavel"
                    else:
                        docx_bytes = modo_fiel_editavel(pdf_file)
                        nome_sufixo = "fiel_editavel_nativo"

                    st.success("Arquivo convertido com sucesso!")

                    st.download_button(
                        label="Baixar Docx",
                        data=docx_bytes,
                        file_name=f"{pdf_file.name.replace('.pdf', '')}_{nome_sufixo}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )
                except Exception as e:
                    st.error(f"Erro ao converter o arquivo: {e}")


if __name__ == "__main__":
    main()

    st.markdown("""
<style>
    .main {background-color: #ffffff; color: #333333;}
    .block-container {padding-top: 1rem; padding-bottom: 0rem;}
    header {display: none !important;}
    footer {display: none !important;}
    #MainMenu {display: none !important;}
    div[data-testid="stAppViewBlockContainer"] {padding-top: 0 !important; padding-bottom: 0 !important;}
    div[data-testid="stVerticalBlock"] {gap: 0 !important; padding-top: 0 !important; padding-bottom: 0 !important;}
    .element-container {margin-top: 0 !important; margin-bottom: 0 !important;}
</style>
""", unsafe_allow_html=True)